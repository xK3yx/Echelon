"""
Text extraction from PDF, DOCX, and plain-text resume files.

Extraction order for PDFs:
  1. pypdf (fast, works for text-based PDFs)
  2. pdfplumber fallback when pypdf either:
       - returns < 100 characters, OR
       - returns "garbled" output where most words are 1-2 chars long
         (graphically-designed PDFs with letter-spacing/kerning often
         produce "P R O G R A M M I N G"-style output via pypdf).
  3. Raise ValueError("scanned_pdf") if pdfplumber also fails.

Raises:
  ValueError("scanned_pdf")     — PDF appears to be scanned / image-only
  Exception                      — any other IO / parsing failure (caller wraps)
"""
from __future__ import annotations

import io

_MIN_TEXT_LENGTH = 100

# Garbled-text detection: if more than this fraction of "words" are 1 char
# long, assume pypdf split each character into its own token (a known
# failure mode for kerned / letter-spaced PDFs) and fall back to pdfplumber.
_GARBLED_SHORT_WORD_THRESHOLD = 0.30
_GARBLED_MIN_WORDS = 50  # only judge once we have enough words to be sure


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    """Dispatch to the correct extractor based on MIME type."""
    if mime_type == "application/pdf":
        return _extract_pdf(file_bytes)
    if mime_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return _extract_docx(file_bytes)
    # text/plain (and any other allowed type)
    return _extract_txt(file_bytes)


# ---------------------------------------------------------------------------
# Internal extractors
# ---------------------------------------------------------------------------


def _looks_garbled(text: str) -> bool:
    """
    Detect pypdf's character-separation failure mode.

    Healthy resume text averages ~5 chars per "word" with very few
    single-character tokens.  Garbled extractions ("P R O G R A M M I N G")
    have most tokens as single characters.
    """
    words = text.split()
    if len(words) < _GARBLED_MIN_WORDS:
        return False
    single_char = sum(1 for w in words if len(w) == 1)
    return (single_char / len(words)) > _GARBLED_SHORT_WORD_THRESHOLD


def _extract_pdf(file_bytes: bytes) -> str:
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()

    needs_fallback = len(text) < _MIN_TEXT_LENGTH or _looks_garbled(text)

    if needs_fallback:
        # Fallback: pdfplumber handles complex layouts and kerned glyphs
        # much more reliably than pypdf at the cost of some speed.
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()

    if len(text) < _MIN_TEXT_LENGTH:
        raise ValueError("scanned_pdf")

    return text


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                stripped = cell.text.strip()
                if stripped:
                    parts.append(stripped)

    return "\n".join(parts)


def _extract_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")
